#!/usr/bin/env python3
"""Generate VPS-3 Auth Stack deployment documentation (DOCX)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "ISIC_Deploiement_VPS3_Auth.docx")


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "1B4F72")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    return table


def build_doc():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    # --- Title page ---
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ISIC — Déploiement VPS-3\nStack Authentification")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("OpenLDAP + Apereo CAS + phpLDAPadmin + Nginx")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"\nInstitut Supérieur de l'Information et de la Communication\n"
        f"Rabat, Maroc\n\n"
        f"Date : {datetime.now().strftime('%d/%m/%Y')}\n"
        f"Version : 1.0\n"
        f"Classification : Confidentiel"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_page_break()

    # --- TOC ---
    doc.add_heading("Table des matières", level=1)
    toc_items = [
        "1. Vue d'ensemble",
        "2. Architecture VPS-3",
        "3. Spécifications serveur",
        "4. Services déployés",
        "5. Configuration OpenLDAP",
        "6. Configuration Apereo CAS",
        "7. Configuration Nginx",
        "8. Sécurité",
        "9. Annuaire LDAP — Structure",
        "10. Utilisateurs et groupes",
        "11. Communication inter-VPS",
        "12. Certificats SSL",
        "13. Procédures d'exploitation",
        "14. Monitoring et logs",
        "15. Plan de reprise",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # --- 1. Vue d'ensemble ---
    doc.add_heading("1. Vue d'ensemble", level=1)
    doc.add_paragraph(
        "Ce document décrit le déploiement du serveur VPS-3, dédié à la stack "
        "d'authentification centralisée de l'ISIC. Ce serveur fournit les services "
        "SSO (Single Sign-On) via le protocole CAS et l'annuaire LDAP pour la "
        "gestion centralisée des identités (étudiants, enseignants, personnel administratif)."
    )
    doc.add_paragraph(
        "Le VPS-3 est un composant critique de l'infrastructure ISIC. Il est utilisé par "
        "le VPS-4 (Odoo prod/staging) pour authentifier les utilisateurs via le protocole CAS "
        "et résoudre les attributs LDAP (uid, mail, rôle, groupe)."
    )

    doc.add_heading("Domaines", level=2)
    add_table(doc,
        ["Domaine", "Service", "Usage"],
        [
            ["auth.isic.ac.ma", "Apereo CAS 6.6.15", "Portail SSO — login centralisé"],
            ["ldap.isic.ac.ma", "phpLDAPadmin", "Administration annuaire LDAP (accès restreint)"],
        ],
        col_widths=[5, 5, 7],
    )

    # --- 2. Architecture ---
    doc.add_heading("2. Architecture VPS-3", level=1)
    doc.add_paragraph(
        "L'architecture suit le modèle Docker Compose avec isolation réseau. "
        "Tous les services communiquent via un réseau interne Docker (bridge, internal: true). "
        "Seul Nginx est exposé sur les ports publics 80 et 443."
    )

    arch_text = """
┌─────────────────────────────────────────────────┐
│                   VPS-3 (51.255.200.172)         │
│                                                   │
│  ┌─────────┐    ┌──────────────┐                 │
│  │  Nginx   │───▶│  CAS 6.6.15  │                │
│  │ (80/443) │    │  (port 8443)  │                │
│  │          │    └──────┬───────┘                 │
│  │          │           │ LDAP                    │
│  │          │    ┌──────▼───────┐    ┌──────────┐│
│  │          │───▶│ phpLDAPadmin │───▶│ OpenLDAP  ││
│  └─────────┘    └──────────────┘    │ (port 389)││
│                                      └─────┬────┘│
│                                            │      │
│  Port 1389 (UFW: VPS-4 only) ◀────────────┘      │
└─────────────────────────────────────────────────┘
         │
         │ LDAP (port 1389)
         ▼
┌─────────────────────┐
│  VPS-4 (51.178.47.204)  │
│  Odoo Prod + Staging     │
└─────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(arch_text.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    # --- 3. Spécifications serveur ---
    doc.add_heading("3. Spécifications serveur", level=1)
    add_table(doc,
        ["Caractéristique", "Valeur"],
        [
            ["Fournisseur", "OVH VPS"],
            ["IP publique", "51.255.200.172"],
            ["OS", "Ubuntu 24.10 (Oracular)"],
            ["RAM", "22 Go"],
            ["CPU", "8 vCPU"],
            ["Disque", "193 Go (191 Go libre)"],
            ["Docker", "28.4.0 + Compose 2.39.2"],
            ["Localisation", "France (OVH)"],
        ],
        col_widths=[6, 11],
    )

    # --- 4. Services déployés ---
    doc.add_heading("4. Services déployés", level=1)
    doc.add_paragraph(
        "Le fichier Docker Compose docker-compose.auth-vps.yml orchestre 4 services :"
    )
    add_table(doc,
        ["Service", "Image", "Container", "Réseau", "RAM max"],
        [
            ["OpenLDAP", "osixia/openldap:1.5.0", "isic-ldap-vps", "interne", "512 Mo"],
            ["phpLDAPadmin", "osixia/phpldapadmin:0.9.0", "isic-phpldapadmin-vps", "interne", "256 Mo"],
            ["CAS", "isic-cas:latest (custom)", "isic-cas-vps", "interne", "2 Go"],
            ["Nginx", "nginx:alpine", "isic-nginx-auth-vps", "interne + externe", "256 Mo"],
        ],
        col_widths=[3, 4.5, 4.5, 2.5, 2.5],
    )

    doc.add_heading("État actuel des services", level=2)
    add_table(doc,
        ["Service", "Status", "Health check"],
        [
            ["isic-ldap-vps", "Running (healthy)", "ldapsearch -x (30s)"],
            ["isic-phpldapadmin-vps", "Running", "—"],
            ["isic-cas-vps", "Running (healthy)", "curl /cas/login (30s, start: 120s)"],
            ["isic-nginx-auth-vps", "Running", "—"],
        ],
        col_widths=[5, 4, 8],
    )

    # --- 5. Configuration OpenLDAP ---
    doc.add_heading("5. Configuration OpenLDAP", level=1)
    doc.add_paragraph(
        "OpenLDAP est configuré avec le domaine isic.ac.ma (base DN: dc=isic,dc=ac,dc=ma). "
        "Les overlays memberOf et refint sont activés pour la gestion automatique "
        "des appartenances aux groupes."
    )

    add_table(doc,
        ["Paramètre", "Valeur"],
        [
            ["Image", "osixia/openldap:1.5.0"],
            ["Base DN", "dc=isic,dc=ac,dc=ma"],
            ["Organisation", "ISIC"],
            ["Backend", "mdb"],
            ["TLS", "Désactivé (communication interne Docker)"],
            ["Port interne", "389"],
            ["Port hôte", "1389 (accès VPS-4 uniquement via UFW)"],
            ["Overlays", "memberOf + refint"],
            ["Admin DN", "cn=admin,dc=isic,dc=ac,dc=ma"],
            ["Volumes", "ldap-data + ldap-config (persistants)"],
        ],
        col_widths=[5, 12],
    )

    doc.add_heading("Overlays activés", level=2)
    doc.add_paragraph(
        "memberOf : Maintient automatiquement l'attribut memberOf sur les entrées "
        "utilisateur lorsqu'elles sont ajoutées/retirées d'un groupe (groupOfNames)."
    )
    doc.add_paragraph(
        "refint (Referential Integrity) : Assure que les suppressions/modifications "
        "de DN sont propagées dans les attributs member, memberOf, manager, owner."
    )

    # --- 6. Configuration CAS ---
    doc.add_heading("6. Configuration Apereo CAS", level=1)
    doc.add_paragraph(
        "CAS 6.6.15 est déployé comme serveur SSO. Il authentifie les utilisateurs "
        "contre l'annuaire LDAP et expose les attributs via le protocole CAS 3.0. "
        "L'authentification statique (casuser/Mellon) est désactivée."
    )

    add_table(doc,
        ["Paramètre", "Valeur"],
        [
            ["Version", "Apereo CAS 6.6.15"],
            ["Image", "isic-cas:latest (custom build)"],
            ["JVM", "-Xms512m -Xmx1536m"],
            ["Port interne", "8443 (HTTP, SSL terminé par Nginx)"],
            ["URL serveur", "https://auth.isic.ac.ma"],
            ["Préfixe", "/cas"],
            ["Auth statique", "Désactivée"],
            ["Auth LDAP", "AUTHENTICATED (bind + search)"],
            ["LDAP URL", "ldap://ldap:389"],
            ["LDAP search base", "ou=People,dc=isic,dc=ac,dc=ma"],
            ["LDAP search filter", "uid={user}"],
            ["Attributs releasés", "uid, cn, mail, sn, givenName, memberOf, employeeType"],
            ["TGT max-time", "28800s (8h)"],
            ["TGT time-to-kill", "7200s (2h)"],
            ["Service registry", "JSON (fichiers /etc/cas/services/)"],
            ["SLO (Single Logout)", "Activé (async)"],
            ["Branding", "Logo ISIC + CSS custom + labels FR"],
        ],
        col_widths=[5, 12],
    )

    doc.add_heading("Service enregistré", level=2)
    doc.add_paragraph(
        "Un service CAS est enregistré pour autoriser toutes les applications "
        "sous le domaine *.isic.ac.ma :"
    )
    add_table(doc,
        ["Champ", "Valeur"],
        [
            ["ID", "1001"],
            ["Nom", "Service d'Authentification Central de l'ISIC"],
            ["Pattern", "^https://.*\\.isic\\.ac\\.ma/.*"],
            ["Attribute Release", "ReturnAllAttributeReleasePolicy"],
        ],
        col_widths=[5, 12],
    )

    # --- 7. Configuration Nginx ---
    doc.add_heading("7. Configuration Nginx", level=1)
    doc.add_paragraph(
        "Nginx sert de reverse proxy SSL pour CAS et phpLDAPadmin. "
        "Il gère la terminaison TLS, la redirection HTTP→HTTPS, et le rate limiting."
    )

    add_table(doc,
        ["Bloc serveur", "Domaine", "Upstream", "Particularités"],
        [
            ["HTTP (80)", "auth + ldap", "—", "Redirect 301 → HTTPS + ACME challenge"],
            ["HTTPS (443)", "auth.isic.ac.ma", "cas:8443", "Rate limit login 10r/m, keepalive 16"],
            ["HTTPS (443)", "ldap.isic.ac.ma", "phpldapadmin:80", "IP restriction (à configurer)"],
        ],
        col_widths=[3.5, 4, 4, 5.5],
    )

    doc.add_heading("Headers de sécurité", level=2)
    doc.add_paragraph("X-Frame-Options: SAMEORIGIN")
    doc.add_paragraph("X-Content-Type-Options: nosniff")
    doc.add_paragraph("X-XSS-Protection: 1; mode=block")
    doc.add_paragraph("Referrer-Policy: strict-origin-when-cross-origin")
    doc.add_paragraph("server_tokens: off (version Nginx masquée)")

    doc.add_heading("HSTS et OCSP (désactivés temporairement)", level=2)
    doc.add_paragraph(
        "Les directives Strict-Transport-Security et ssl_stapling sont commentées "
        "dans la configuration actuelle car les certificats SSL sont auto-signés. "
        "Elles seront réactivées lors de la mise en place de Let's Encrypt avec "
        "les certificats de production."
    )

    # --- 8. Sécurité ---
    doc.add_heading("8. Sécurité", level=1)

    doc.add_heading("Pare-feu UFW", level=2)
    add_table(doc,
        ["Règle", "Port", "Source", "Description"],
        [
            ["ALLOW", "22/tcp", "Anywhere", "SSH"],
            ["ALLOW", "80/tcp", "Anywhere", "HTTP (redirect HTTPS)"],
            ["ALLOW", "443/tcp", "Anywhere", "HTTPS (CAS + phpLDAPadmin)"],
            ["ALLOW", "1389/tcp", "51.178.47.204", "LDAP (VPS-4 uniquement)"],
            ["DENY", "Tout le reste", "Anywhere", "Politique par défaut"],
        ],
        col_widths=[2.5, 3, 4, 7.5],
    )

    doc.add_heading("fail2ban", level=2)
    add_table(doc,
        ["Paramètre", "Valeur"],
        [
            ["Jail", "sshd"],
            ["Max retries", "3"],
            ["Find time", "600s (10 min)"],
            ["Ban time", "3600s (1h)"],
            ["Action", "UFW ban"],
            ["Log", "/var/log/auth.log"],
        ],
        col_widths=[6, 11],
    )

    doc.add_heading("Isolation réseau Docker", level=2)
    doc.add_paragraph(
        "Les services LDAP, CAS et phpLDAPadmin communiquent via le réseau "
        "isic-auth-internal-vps (bridge, internal: true). Ce réseau n'a pas "
        "d'accès Internet sortant. Seul Nginx est connecté au réseau externe "
        "(isic-auth-external-vps) pour recevoir le trafic entrant."
    )

    doc.add_heading("Restriction IP phpLDAPadmin", level=2)
    doc.add_paragraph(
        "IMPORTANT : La configuration Nginx inclut un bloc commenté pour restreindre "
        "l'accès à phpLDAPadmin (ldap.isic.ac.ma) par IP. Avant la mise en production, "
        "décommenter et configurer les IPs autorisées :"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "allow 196.200.xxx.xxx;    # Réseau campus ISIC\n"
        "allow 105.xxx.xxx.xxx;    # IP admin\n"
        "deny all;"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # --- 9. Structure LDAP ---
    doc.add_heading("9. Annuaire LDAP — Structure", level=1)
    doc.add_paragraph("L'arborescence LDAP est organisée comme suit :")

    p = doc.add_paragraph()
    tree = (
        "dc=isic,dc=ac,dc=ma\n"
        "├── ou=People\n"
        "│   ├── ou=Staff          (personnel ISIC)\n"
        "│   └── ou=Students       (étudiants)\n"
        "├── ou=Groups\n"
        "│   ├── cn=direction      (groupe direction)\n"
        "│   ├── cn=enseignants    (groupe enseignants)\n"
        "│   ├── cn=scolarite      (groupe scolarité)\n"
        "│   └── cn=etudiants      (groupe étudiants)\n"
        "└── ou=Services\n"
        "    └── uid=odoo_bind     (compte technique Odoo)"
    )
    run = p.add_run(tree)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # --- 10. Utilisateurs et groupes ---
    doc.add_heading("10. Utilisateurs et groupes", level=1)

    doc.add_heading("Personnel (ou=Staff)", level=2)
    add_table(doc,
        ["UID", "Nom complet", "Rôle", "Groupe"],
        [
            ["directeur", "Mohammed Directeur", "Directeur ISIC", "direction"],
            ["admin.scolarite", "Fatima Admin", "Chef service scolarité", "scolarite"],
            ["agent.scolarite", "Karim Agent", "Agent scolarité", "scolarite"],
            ["resp.ci", "Ahmed Responsable", "Resp. Communication", "enseignants"],
            ["prof.informatique", "Sara Prof", "Enseignante informatique", "enseignants"],
        ],
        col_widths=[4, 4, 4.5, 4.5],
    )

    doc.add_heading("Étudiants (ou=Students)", level=2)
    add_table(doc,
        ["UID", "Nom complet", "Groupe"],
        [
            ["etudiant1", "Youssef Etudiant", "etudiants"],
            ["etudiant2", "Zineb Etudiante", "etudiants"],
            ["etudiant3", "Omar Etudiant", "etudiants"],
        ],
        col_widths=[5, 7, 5],
    )

    doc.add_heading("Compte technique", level=2)
    doc.add_paragraph(
        "uid=odoo_bind (ou=Services) : compte de service utilisé par Odoo "
        "pour interroger l'annuaire LDAP. Permissions en lecture seule sur "
        "l'ensemble de l'arborescence."
    )

    # --- 11. Communication inter-VPS ---
    doc.add_heading("11. Communication inter-VPS", level=1)
    doc.add_paragraph(
        "Le VPS-4 (Odoo) communique avec le VPS-3 via le port LDAP 1389 "
        "pour l'authentification CAS et la résolution d'attributs."
    )

    add_table(doc,
        ["Flux", "Source", "Destination", "Port", "Protocole"],
        [
            ["Auth CAS", "Navigateur → VPS-3", "Nginx → CAS", "443", "HTTPS (CAS 3.0)"],
            ["Ticket CAS", "Odoo (VPS-4) → VPS-3", "CAS", "443", "HTTPS (validation)"],
            ["LDAP", "CAS (interne)", "OpenLDAP", "389", "LDAP (bind+search)"],
            ["LDAP externe", "VPS-4 (51.178.47.204)", "OpenLDAP", "1389", "LDAP (UFW filtered)"],
        ],
        col_widths=[3, 4, 3.5, 2, 4.5],
    )

    doc.add_heading("Flux CAS complet", level=2)
    p = doc.add_paragraph()
    flow = (
        "1. Utilisateur accède à Odoo (mon.isic.ac.ma)\n"
        "2. Odoo redirige vers auth.isic.ac.ma/cas/login?service=...\n"
        "3. CAS affiche la page de login ISIC\n"
        "4. Utilisateur saisit uid + mot de passe LDAP\n"
        "5. CAS authentifie via LDAP (bind + search)\n"
        "6. CAS redirige vers Odoo avec un ticket ST-xxx\n"
        "7. Odoo valide le ticket auprès de CAS (serviceValidate)\n"
        "8. CAS retourne les attributs (uid, mail, memberOf...)\n"
        "9. Odoo crée/lie l'utilisateur et ouvre la session"
    )
    run = p.add_run(flow)
    run.font.size = Pt(9)

    # --- 12. Certificats SSL ---
    doc.add_heading("12. Certificats SSL", level=1)
    doc.add_paragraph(
        "Des certificats auto-signés sont actuellement en place pour permettre "
        "le fonctionnement HTTPS. Ils seront remplacés par des certificats "
        "Let's Encrypt lorsque les enregistrements DNS publics seront configurés."
    )

    add_table(doc,
        ["Fichier", "Emplacement", "Description"],
        [
            ["fullchain.pem", "/opt/isic/docker/nginx/ssl/", "Certificat auto-signé (365 jours)"],
            ["privkey.pem", "/opt/isic/docker/nginx/ssl/", "Clé privée RSA 2048"],
            ["chain.pem", "/opt/isic/docker/nginx/ssl/", "Copie fullchain (OCSP placeholder)"],
        ],
        col_widths=[4, 7, 6],
    )

    doc.add_heading("SAN (Subject Alternative Names)", level=2)
    doc.add_paragraph("DNS: auth.isic.ac.ma, ldap.isic.ac.ma")

    doc.add_heading("Migration Let's Encrypt", level=2)
    doc.add_paragraph(
        "Pré-requis : enregistrements DNS A pour auth.isic.ac.ma et ldap.isic.ac.ma "
        "pointant vers 51.255.200.172. Une fois les DNS propagés, utiliser certbot "
        "pour obtenir des certificats valides et réactiver HSTS + OCSP stapling."
    )

    # --- 13. Procédures d'exploitation ---
    doc.add_heading("13. Procédures d'exploitation", level=1)

    doc.add_heading("Démarrage complet", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "cd /opt/isic/docker\n"
        "docker compose -f docker-compose.auth-vps.yml --env-file .env.vps3 up -d"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Arrêt complet", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "cd /opt/isic/docker\n"
        "docker compose -f docker-compose.auth-vps.yml --env-file .env.vps3 down"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Redémarrage d'un service", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "docker restart isic-cas-vps       # Redémarrer CAS\n"
        "docker restart isic-ldap-vps      # Redémarrer LDAP\n"
        "docker restart isic-nginx-auth-vps  # Redémarrer Nginx"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Vérification santé des services", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'docker ps --format "table {{.Names}}\\t{{.Status}}"\n'
        "docker exec isic-ldap-vps ldapsearch -x -H ldap://localhost:389 "
        '-b "dc=isic,dc=ac,dc=ma" -D "cn=admin,dc=isic,dc=ac,dc=ma" -w $LDAP_ADMIN_PASSWORD\n'
        "curl -sk https://localhost/cas/login | grep '<title>'"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Ajout d'un utilisateur LDAP", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "# Créer un fichier new_user.ldif, puis :\n"
        "docker cp new_user.ldif isic-ldap-vps:/tmp/\n"
        "docker exec isic-ldap-vps ldapadd -x -H ldap://localhost:389 \\\n"
        '  -D "cn=admin,dc=isic,dc=ac,dc=ma" -w $LDAP_ADMIN_PASSWORD \\\n'
        "  -f /tmp/new_user.ldif\n"
        "\n"
        "# Ou via phpLDAPadmin : https://ldap.isic.ac.ma"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Mise à jour de l'image CAS", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "cd /opt/isic/docker\n"
        "docker compose -f docker-compose.auth-vps.yml --env-file .env.vps3 build cas\n"
        "docker compose -f docker-compose.auth-vps.yml --env-file .env.vps3 up -d cas"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # --- 14. Monitoring et logs ---
    doc.add_heading("14. Monitoring et logs", level=1)

    add_table(doc,
        ["Service", "Commande logs", "Rotation"],
        [
            ["LDAP", "docker logs isic-ldap-vps", "json-file 10m × 3"],
            ["CAS", "docker logs isic-cas-vps", "json-file 10m × 3"],
            ["phpLDAPadmin", "docker logs isic-phpldapadmin-vps", "json-file 10m × 3"],
            ["Nginx", "docker logs isic-nginx-auth-vps", "json-file 10m × 3"],
            ["fail2ban", "sudo fail2ban-client status sshd", "syslog"],
        ],
        col_widths=[4, 7, 6],
    )

    doc.add_heading("Vérification fail2ban", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "sudo fail2ban-client status sshd      # IPs bannies\n"
        "sudo fail2ban-client set sshd unbanip 1.2.3.4  # Débannir une IP"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # --- 15. Plan de reprise ---
    doc.add_heading("15. Plan de reprise", level=1)

    doc.add_heading("Données persistantes", level=2)
    add_table(doc,
        ["Volume Docker", "Contenu", "Criticité"],
        [
            ["isic-ldap-data-vps", "Base de données LDAP (mdb)", "Critique"],
            ["isic-ldap-config-vps", "Configuration LDAP (slapd.d)", "Critique"],
            ["isic-cas-logs-vps", "Logs CAS", "Faible"],
            ["isic-certbot-auth-vps", "Challenges ACME (Let's Encrypt)", "Faible"],
        ],
        col_widths=[5, 7, 5],
    )

    doc.add_heading("Sauvegarde LDAP", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "# Export complet de l'annuaire LDAP (LDIF)\n"
        "docker exec isic-ldap-vps slapcat > /opt/isic/backup/ldap-$(date +%Y%m%d).ldif\n"
        "\n"
        "# Restauration\n"
        "docker cp backup.ldif isic-ldap-vps:/tmp/\n"
        "docker exec isic-ldap-vps slapadd -l /tmp/backup.ldif"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("Reconstruction complète", level=2)
    doc.add_paragraph(
        "En cas de perte totale du serveur, les fichiers d'infrastructure sont "
        "versionnés dans le dépôt Git isic-v1. La procédure de reconstruction est :"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "1. Provisionner un nouveau VPS Ubuntu\n"
        "2. Installer Docker + UFW + fail2ban\n"
        "3. Cloner le dépôt isic-v1\n"
        "4. Créer .env.vps3 depuis le template\n"
        "5. Générer les certificats SSL\n"
        "6. docker compose up -d\n"
        "7. Exécuter init-ldap-vps.sh\n"
        "8. Restaurer le dump LDAP si disponible"
    )
    run.font.size = Pt(9)

    # --- Footer ---
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "— Fin du document —\n\n"
        "Document généré automatiquement\n"
        f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        "ISIC — Institut Supérieur de l'Information et de la Communication"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Document saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_doc()
